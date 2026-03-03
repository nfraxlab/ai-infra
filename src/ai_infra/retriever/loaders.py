"""File loaders for the Retriever module.

Loads various file formats and extracts text content with metadata.
Uses LangChain loaders internally where available.
"""

from __future__ import annotations

import json
import logging
import os
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Type alias for loaded documents: (text, metadata)
LoadedDocument = tuple[str, dict[str, Any]]


def load_file(path: str) -> list[LoadedDocument]:
    """Load a file and return its text content with metadata.

    Automatically detects the file type and uses the appropriate loader.

    Args:
        path: Path to the file to load.

    Returns:
        List of (text, metadata) tuples. Some formats (like PDF) may
        return multiple documents (one per page).

    Raises:
        FileNotFoundError: If the file doesn't exist.
        ValueError: If the file type is not supported.
    """
    path = os.path.expanduser(path)
    if not os.path.isfile(path):
        raise FileNotFoundError(f"File not found: {path}")

    ext = Path(path).suffix.lower()

    loaders = {
        ".pdf": load_pdf,
        ".txt": load_text,
        ".md": load_text,
        ".markdown": load_text,
        ".docx": load_docx,
        ".doc": load_docx,
        ".csv": load_csv,
        ".json": load_json,
        ".html": load_html,
        ".htm": load_html,
        ".xlsx": load_excel,
        ".xls": load_excel,
    }

    loader = loaders.get(ext)
    if loader is None:
        supported = ", ".join(sorted(loaders.keys()))
        raise ValueError(f"Unsupported file type: {ext}\nSupported types: {supported}")

    return loader(path)


def load_directory(
    path: str,
    pattern: str = "*",
    recursive: bool = True,
) -> list[LoadedDocument]:
    """Load all files from a directory.

    Args:
        path: Path to the directory.
        pattern: Glob pattern for file matching (e.g., "*.pdf", "*.md").
        recursive: Whether to search subdirectories.

    Returns:
        List of (text, metadata) tuples from all loaded files.

    Raises:
        FileNotFoundError: If the directory doesn't exist.
    """
    path = os.path.expanduser(path)
    if not os.path.isdir(path):
        raise FileNotFoundError(f"Directory not found: {path}")

    dir_path = Path(path)
    all_documents: list[LoadedDocument] = []

    # Get matching files
    if recursive:
        files = list(dir_path.rglob(pattern))
    else:
        files = list(dir_path.glob(pattern))

    # Filter to only files (not directories)
    files = [f for f in files if f.is_file()]

    # Load each file
    for file_path in sorted(files):
        try:
            docs = load_file(str(file_path))
            all_documents.extend(docs)
        except ValueError:
            # Skip unsupported file types
            continue

    return all_documents


# ==============================================================================
# Individual file type loaders
# ==============================================================================


def load_pdf(path: str) -> list[LoadedDocument]:
    """Load a PDF file, returning one document per page.

    Uses pypdf for digital PDFs. When a page yields no text,
    falls back to OCR via pypdfium2 + pytesseract (if installed).

    Args:
        path: Path to the PDF file.

    Returns:
        List of (text, metadata) tuples, one per page.
    """
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError(
            "pypdf is required for PDF loading. Install it with: pip install pypdf"
        ) from e

    documents: list[LoadedDocument] = []
    reader = PdfReader(path)
    total_pages = len(reader.pages)
    ocr_pages: list[int] = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            metadata = {
                "source": path,
                "page": i + 1,
                "total_pages": total_pages,
            }
            documents.append((text, metadata))
        else:
            # Page has no text — attempt OCR fallback
            ocr_text = _ocr_pdf_page(path, i)
            if ocr_text.strip():
                ocr_pages.append(i + 1)
                metadata = {
                    "source": path,
                    "page": i + 1,
                    "total_pages": total_pages,
                    "ocr": True,
                }
                documents.append((ocr_text, metadata))

    if ocr_pages:
        logger.info("pdf_ocr_fallback", extra={"path": path, "ocr_pages": ocr_pages})

    return documents


def _ocr_pdf_page(path: str, page_index: int, dpi: int = 200) -> str:
    """OCR a single page of a PDF using pypdfium2 + pytesseract.

    Uses pypdfium2 to render the page to a PIL image (no poppler needed),
    then pytesseract for text recognition.

    Falls back gracefully if dependencies are not installed.

    Args:
        path: Path to the PDF file.
        page_index: Zero-based page index.
        dpi: Resolution for rendering. Default 200.

    Returns:
        Extracted text, or empty string if OCR is unavailable.
    """
    try:
        import pypdfium2 as pdfium
        import pytesseract
    except ImportError:
        logger.debug(
            "ocr_unavailable",
            extra={
                "reason": "pypdfium2 and/or pytesseract not installed",
                "hint": "pip install 'ai-infra[ocr]'",
            },
        )
        return ""

    try:
        pdf = pdfium.PdfDocument(path)
        page = pdf[page_index]
        bitmap = page.render(scale=dpi / 72)
        img = bitmap.to_pil()
        text: str = pytesseract.image_to_string(img, config="--psm 6")
        return text.strip()
    except Exception:
        logger.warning(
            "ocr_page_failed",
            extra={"path": path, "page": page_index + 1},
            exc_info=True,
        )
        return ""


def ocr_page_from_bytes(data: bytes, page_number: int, dpi: int = 200) -> str:
    """OCR a single page of a PDF from raw bytes.

    Public API for callers that work with in-memory PDF bytes
    (e.g. file uploads) rather than file paths.

    Uses pypdfium2 to render the page to a PIL image (no poppler needed),
    then pytesseract for text recognition.

    Args:
        data: Raw PDF bytes.
        page_number: 1-based page number to OCR.
        dpi: Resolution for rendering. Default 200.

    Returns:
        Extracted text from the page, or empty string if OCR
        is unavailable or fails.
    """
    try:
        import pypdfium2 as pdfium
        import pytesseract
    except ImportError:
        logger.debug(
            "ocr_unavailable",
            extra={
                "reason": "pypdfium2 and/or pytesseract not installed",
                "hint": "pip install 'ai-infra[ocr]'",
            },
        )
        return ""

    try:
        pdf = pdfium.PdfDocument(data)
        page = pdf[page_number - 1]  # pypdfium2 is 0-indexed
        bitmap = page.render(scale=dpi / 72)
        img = bitmap.to_pil()
        text: str = pytesseract.image_to_string(img, config="--psm 6")
        return text.strip()
    except Exception:
        logger.warning(
            "ocr_page_failed",
            extra={"page": page_number},
            exc_info=True,
        )
        return ""


def load_text(path: str) -> list[LoadedDocument]:
    """Load a plain text or markdown file.

    Args:
        path: Path to the text file.

    Returns:
        List containing a single (text, metadata) tuple.
    """
    with open(path, encoding="utf-8") as f:
        text = f.read()

    metadata = {
        "source": path,
        "file_type": Path(path).suffix.lower(),
    }
    return [(text, metadata)]


def load_docx(path: str) -> list[LoadedDocument]:
    """Load a DOCX file.

    Args:
        path: Path to the DOCX file.

    Returns:
        List containing a single (text, metadata) tuple.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX loading. Install it with: pip install python-docx"
        ) from e

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    metadata = {
        "source": path,
        "file_type": ".docx",
    }
    return [(text, metadata)]


def load_csv(path: str) -> list[LoadedDocument]:
    """Load a CSV file, converting rows to text.

    Each row is converted to a text representation with column names.

    Args:
        path: Path to the CSV file.

    Returns:
        List of (text, metadata) tuples, one per row.
    """
    try:
        import pandas as pd
    except ImportError as e:
        raise ImportError(
            "pandas is required for CSV loading. Install it with: pip install pandas"
        ) from e

    df = pd.read_csv(path)
    documents: list[LoadedDocument] = []

    for i, row in df.iterrows():
        # Convert row to readable text format
        parts = [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
        text = "\n".join(parts)

        if text.strip():
            metadata = {
                "source": path,
                "row_index": int(i) if isinstance(i, int) else str(i),
                "file_type": ".csv",
            }
            documents.append((text, metadata))

    return documents


def load_json(path: str) -> list[LoadedDocument]:
    """Load a JSON file.

    Handles both JSON objects and JSON arrays.

    Args:
        path: Path to the JSON file.

    Returns:
        List of (text, metadata) tuples.
    """
    with open(path, encoding="utf-8") as f:
        data = json.load(f)

    documents: list[LoadedDocument] = []

    if isinstance(data, list):
        # JSON array - one document per item
        for i, item in enumerate(data):
            text = json.dumps(item, indent=2, ensure_ascii=False)
            metadata = {
                "source": path,
                "item_index": i,
                "file_type": ".json",
            }
            documents.append((text, metadata))
    else:
        # Single JSON object
        text = json.dumps(data, indent=2, ensure_ascii=False)
        metadata = {
            "source": path,
            "file_type": ".json",
        }
        documents.append((text, metadata))

    return documents


def load_html(path: str) -> list[LoadedDocument]:
    """Load an HTML file, extracting text content.

    Args:
        path: Path to the HTML file.

    Returns:
        List containing a single (text, metadata) tuple.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError as e:
        raise ImportError(
            "beautifulsoup4 is required for HTML loading. "
            "Install it with: pip install beautifulsoup4"
        ) from e

    with open(path, encoding="utf-8") as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, "html.parser")

    # Remove script and style elements
    for script in soup(["script", "style"]):
        script.decompose()

    # Get text
    text = soup.get_text(separator="\n")

    # Clean up whitespace
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(line for line in lines if line)

    # Try to extract title
    title = None
    if soup.title and soup.title.string:
        title = soup.title.string.strip()

    metadata = {
        "source": path,
        "file_type": ".html",
    }
    if title:
        metadata["title"] = title

    return [(text, metadata)]


def load_excel(path: str) -> list[LoadedDocument]:
    """Load an Excel file (.xlsx/.xls), returning one document per sheet.

    Each row is converted to a pipe-delimited text representation
    with column headers.

    Args:
        path: Path to the Excel file.

    Returns:
        List of (text, metadata) tuples, one per non-empty sheet.

    Raises:
        ImportError: If openpyxl is not installed.
    """
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise ImportError(
            "openpyxl is required for Excel loading. "
            "Install it with: pip install openpyxl "
            "or pip install 'ai-infra[excel]'"
        ) from e

    wb = load_workbook(path, read_only=True, data_only=True)
    documents: list[LoadedDocument] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = [r for r in ws.iter_rows(values_only=True) if not all(cell is None for cell in r)]
        if not rows:
            continue

        # First row as headers
        headers = [str(h) if h is not None else "" for h in rows[0]]
        lines = [" | ".join(headers)]

        for row in rows[1:]:
            cells = [str(c) if c is not None else "" for c in row]
            lines.append(" | ".join(cells))

        text = "\n".join(lines)
        if text.strip():
            metadata = {
                "source": path,
                "sheet": sheet_name,
                "file_type": Path(path).suffix.lower(),
            }
            documents.append((text, metadata))

    wb.close()
    return documents


def load_excel_bytes(data: bytes) -> list[LoadedDocument]:
    """Load an Excel file from raw bytes, returning one document per sheet.

    Public API for callers that work with in-memory bytes
    (e.g. file uploads) rather than file paths.

    Args:
        data: Raw .xlsx/.xls bytes.

    Returns:
        List of (text, metadata) tuples, one per non-empty sheet.

    Raises:
        ImportError: If openpyxl is not installed.
        ValueError: If the workbook has no data.
    """
    import io as _io

    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise ImportError(
            "openpyxl is required for Excel loading. "
            "Install it with: pip install openpyxl "
            "or pip install 'ai-infra[excel]'"
        ) from e

    wb = load_workbook(_io.BytesIO(data), read_only=True, data_only=True)
    documents: list[LoadedDocument] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = [r for r in ws.iter_rows(values_only=True) if not all(cell is None for cell in r)]
        if not rows:
            continue

        headers = [str(h) if h is not None else "" for h in rows[0]]
        lines = [" | ".join(headers)]

        for row in rows[1:]:
            cells = [str(c) if c is not None else "" for c in row]
            lines.append(" | ".join(cells))

        text = "\n".join(lines)
        if text.strip():
            metadata = {
                "source": "<bytes>",
                "sheet": sheet_name,
                "file_type": ".xlsx",
            }
            documents.append((text, metadata))

    wb.close()

    if not documents:
        raise ValueError("The Excel file contains no data.")

    return documents
